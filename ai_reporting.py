import os
import google.generativeai as genai
from docx import Document
from docx.shared import Inches
from fpdf import FPDF
import pandas as pd
import io
import streamlit as st

class AIReportingLayer:
    def __init__(self, api_key):
        self.api_key = api_key
        if api_key:
            genai.configure(api_key=api_key)
            # Comprehensive list for robust fallback
            self.model_names = [
                'gemini-2.0-flash', 
                'gemini-2.0-flash-lite',
                'gemini-2.5-flash', 
                'gemini-3-flash-preview',
                'gemini-2.0-flash-001',
                'gemini-2.0-flash-lite-001',
                'gemini-2.0-flash-lite-preview',
                'gemini-1.5-flash', 
                'gemini-1.5-flash-8b',
                'gemini-1.5-flash-latest',
                'gemini-1.5-flash-lite-latest',
                'gemini-2.5-pro',
                'gemini-3-pro-preview',
                'gemini-2.0-pro-exp',
                'gemini-1.5-pro',
                'gemini-1.5-pro-latest',
                'gemini-1.5-pro-002',
                'gemini-exp-1206',
                'gemini-pro-latest',
                'gemma-3-27b-it',
                'gemma-3-12b-it',
                'gemma-3-4b-it'
            ]
            self.model = None
            for name in self.model_names:
                try:
                    self.model = genai.GenerativeModel(name)
                    # Test if it works with a simple query if needed, 
                    # but usually initialization is enough to catch simple name errors.
                    break
                except:
                    continue
        else:
            self.model = None

    def generate_analysis(self, strategy_name, ticker, stats, signals_summary):
        """Generate AI strategy analysis, market sentiment, and risk assessment."""
        if not self.model:
            return "API Key not configured. Please add GOOGLE_API_KEY to secrets.toml."

        prompt = f"""
        Act as an institutional quantitative researcher. Analyze the following backtest results for {ticker} using the {strategy_name} strategy.
        
        ### DATA INPUTS ###
        Backtest Statistics:
        {stats.to_string()}
        
        Recent Trade Signals:
        {signals_summary}
        
        ### REQUIRED REPORT STRUCTURE ###
        You MUST provide the analysis in the following structure:
        
        1. **SENTIMENT_TAG**: [BULLISH / BEARISH / NEUTRAL] (One word only)
        
        2. **Strategy Analysis**:
           - Evaluate the core efficiency of the {strategy_name} logic.
           - Discuss the win/loss distribution based on the total trades.
        
        3. **Market Sentiment**:
           - Analyze the immediate price action trend.
           - Determine if the current technical setup favors further expansion.
        
        4. **Risk Assessment**:
           - Deep dive into the Maximum Drawdown ({stats.get('Max Drawdown', 'N/A')}).
           - Evaluate the Sharpe Ratio ({stats.get('Sharpe Ratio', 'N/A')}) for risk-adjusted returns.
           - Mention specific tail risks.
        
        5. **Investment Conclusion**:
           - Final institutional-grade summary.

        Format the entire response in clean Markdown.
        """
        
        try:
            # Try recursive fallback if one model fails (not found or quota)
            for name in self.model_names:
                try:
                    current_model = genai.GenerativeModel(name)
                    response = current_model.generate_content(prompt)
                    return response.text
                except Exception as inner_e:
                    err_msg = str(inner_e).lower()
                    
                    # If it's a 429 (Quota Exceeded), STOP. 
                    # Looping through 18 models on a 429 just burns your Daily RPD limit 18x faster.
                    if "429" in err_msg or "quota" in err_msg:
                        return f"🛑 Gemini Quota Exceeded (429). The Google Free Tier has reached its limit. Model tried: {name}. Please wait 60 seconds or check https://aistudio.google.com/."
                        
                    # Only fallback for 404 (not found)
                    if "404" in err_msg or "not found" in err_msg:
                        continue
                        
                    raise inner_e
            return "⚠️ All available Gemini models are currently offline or unsupported by your API key."
        except Exception as e:
            return f"Error generating AI analysis: {str(e)}"

    def extract_sentiment(self, ai_text):
        """Extract the sentiment tag from the AI response."""
        if "BULLISH" in ai_text.upper(): return "BULLISH"
        if "BEARISH" in ai_text.upper(): return "BEARISH"
        return "NEUTRAL"

    def create_word_report(self, strategy_name, ticker, stats, ai_analysis):
        """Generate a professional Word report."""
        doc = Document()
        doc.add_heading(f'ZMTech Quant Analysis Report: {ticker}', 0)
        
        doc.add_heading('Strategy Overview', level=1)
        doc.add_paragraph(f'Strategy: {strategy_name}')
        doc.add_paragraph(f'Ticker: {ticker}')
        
        doc.add_heading('Performance Metrics', level=1)
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Value'
        
        for k, v in stats.items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(k)
            row_cells[1].text = f"{v:.2f}" if isinstance(v, float) else str(v)
            
        doc.add_heading('AI Intelligence & Analysis', level=1)
        doc.add_paragraph(ai_analysis)
        
        doc.add_heading('Disclaimer', level=1)
        doc.add_paragraph('This report is for informational purposes only and does not constitute financial advice.')
        
        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()

    def create_pdf_report(self, strategy_name, ticker, stats, ai_analysis):
        """Generate a professional PDF report with robust encoding handling."""
        try:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("helvetica", 'B', 16)
            pdf.cell(0, 10, f'ZMTech Quant Analysis Report: {ticker}', ln=True, align='C')
            pdf.ln(10)
            
            pdf.set_font("helvetica", 'B', 12)
            pdf.cell(0, 10, 'Strategy Overview', ln=True)
            pdf.set_font("helvetica", '', 10)
            pdf.cell(0, 10, f'Strategy: {strategy_name}', ln=True)
            pdf.cell(0, 10, f'Ticker: {ticker}', ln=True)
            pdf.ln(5)
            
            pdf.set_font("helvetica", 'B', 12)
            pdf.cell(0, 10, 'Performance Metrics', ln=True)
            pdf.set_font("helvetica", '', 10)
            for k, v in stats.items():
                val_str = f"{v:.4f}" if isinstance(v, float) else str(v)
                pdf.cell(0, 8, f'{k}: {val_str}', ln=True)
            pdf.ln(5)
            
            pdf.set_font("helvetica", 'B', 12)
            pdf.cell(0, 10, 'AI Intelligence & Analysis', ln=True)
            pdf.set_font("helvetica", '', 10)
            
            # Robust text cleaning for PDF
            clean_text = ai_analysis.replace('**', '').replace('#', '').replace('`', '').replace('*', '-')
            # Ensure latin-1 compatibility for standard FPDF fonts
            clean_text = clean_text.encode('latin-1', 'replace').decode('latin-1')
            
            pdf.multi_cell(0, 8, clean_text)
            
            pdf.ln(10)
            pdf.set_font("helvetica", 'I', 8)
            pdf.cell(0, 10, 'Disclaimer: For informational purposes only. No financial advice.', ln=True, align='C')
            
            return bytes(pdf.output())
        except Exception as e:
            st.error(f"PDF Generation Error: {str(e)}")
            return b""
